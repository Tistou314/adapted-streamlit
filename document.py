"""
Document extraction (Word/PDF) and reconstruction for AdaptEd.

Extracts text and images in reading order from .docx and .pdf files,
builds marked text with figure placeholders for Claude,
and reconstructs documents with adapted text + original images.
"""

import base64
import io
import re
from dataclasses import dataclass, field
from typing import Literal

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.shared import Pt, Inches
from PIL import Image


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class ContentBlock:
    """A single block in the document's reading order."""
    type: Literal["text", "image"]
    text: str = ""
    image_bytes: bytes = b""
    image_mime: str = "image/png"
    width_pt: float | None = None
    height_pt: float | None = None
    image_index: int = 0


FIGURE_MARKER = "{{{{FIGURE_ORIGINALE_{n}}}}}"


# ---------------------------------------------------------------------------
# DOCX extraction
# ---------------------------------------------------------------------------

# XML namespaces used in .docx files
_NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
}

EMU_PER_PT = 914400 / 72  # 1 point = 12700 EMU


def extract_docx_content(file_bytes: bytes) -> list[ContentBlock]:
    """Extract text and images from a .docx file in document order."""
    doc = DocxDocument(io.BytesIO(file_bytes))
    blocks: list[ContentBlock] = []
    image_counter = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Check if paragraph contains images
            blips = element.findall(".//a:blip", _NSMAP)
            if blips:
                # Extract text before/around images (if any)
                para_text = _get_paragraph_text(element)
                if para_text.strip():
                    blocks.append(ContentBlock(type="text", text=para_text.strip()))

                # Extract each image
                for blip in blips:
                    r_embed = blip.get(
                        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                    )
                    if r_embed:
                        try:
                            image_part = doc.part.related_parts[r_embed]
                            img_bytes = image_part.blob
                            mime = image_part.content_type or "image/png"

                            # Try to get dimensions from wp:extent
                            width_pt, height_pt = _get_image_dimensions(element)

                            blocks.append(ContentBlock(
                                type="image",
                                image_bytes=img_bytes,
                                image_mime=mime,
                                width_pt=width_pt,
                                height_pt=height_pt,
                                image_index=image_counter,
                            ))
                            image_counter += 1
                        except (KeyError, AttributeError):
                            pass
            else:
                # Text-only paragraph
                para_text = _get_paragraph_text(element)
                if para_text.strip():
                    blocks.append(ContentBlock(type="text", text=para_text.strip()))

        elif tag == "tbl":
            # Convert table to markdown-like text
            table_text = _table_to_text(element)
            if table_text.strip():
                blocks.append(ContentBlock(type="text", text=table_text))

    return _consolidate_text_blocks(blocks)


def _get_paragraph_text(para_element) -> str:
    """Extract text from a w:p element."""
    texts = []
    for t_elem in para_element.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t"):
        if t_elem.text:
            texts.append(t_elem.text)
    return "".join(texts)


def _get_image_dimensions(para_element) -> tuple[float | None, float | None]:
    """Extract image dimensions from wp:extent in EMUs, convert to points."""
    extents = para_element.findall(".//wp:extent", _NSMAP)
    if extents:
        extent = extents[0]
        cx = extent.get("cx")
        cy = extent.get("cy")
        if cx and cy:
            width_pt = int(cx) / EMU_PER_PT
            height_pt = int(cy) / EMU_PER_PT
            return width_pt, height_pt
    return None, None


def _table_to_text(tbl_element) -> str:
    """Convert a w:tbl element to a simple markdown table."""
    rows = tbl_element.findall(".//w:tr", {"w": _NSMAP["w"]})
    if not rows:
        return ""

    table_rows: list[list[str]] = []
    for row in rows:
        cells = row.findall(".//w:tc", {"w": _NSMAP["w"]})
        row_texts = []
        for cell in cells:
            cell_text = ""
            for p in cell.findall(".//w:p", {"w": _NSMAP["w"]}):
                cell_text += _get_paragraph_text(p) + " "
            row_texts.append(cell_text.strip())
        table_rows.append(row_texts)

    if not table_rows:
        return ""

    # Build markdown table
    lines = []
    lines.append("| " + " | ".join(table_rows[0]) + " |")
    lines.append("| " + " | ".join(["---"] * len(table_rows[0])) + " |")
    for row in table_rows[1:]:
        # Pad row if needed
        while len(row) < len(table_rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


# ---------------------------------------------------------------------------
# PDF extraction
# ---------------------------------------------------------------------------

def extract_pdf_content(file_bytes: bytes) -> list[ContentBlock]:
    """Extract text and images from a PDF file in reading order."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    blocks: list[ContentBlock] = []
    image_counter = 0

    for page in doc:
        page_dict = page.get_text("dict")

        # Also get image references for reliable byte extraction
        page_images = page.get_images(full=True)
        # Map xref -> extracted PNG bytes
        xref_to_bytes: dict[int, bytes] = {}

        for block in page_dict["blocks"]:
            if block["type"] == 0:  # Text block
                text = ""
                for line in block["lines"]:
                    line_text = "".join(span["text"] for span in line["spans"])
                    text += line_text + "\n"
                if text.strip():
                    blocks.append(ContentBlock(type="text", text=text.strip()))

            elif block["type"] == 1:  # Image block
                bbox = block["bbox"]
                width_pt = bbox[2] - bbox[0]
                height_pt = bbox[3] - bbox[1]

                # Skip tiny images (likely decorations/bullets)
                if width_pt < 30 or height_pt < 30:
                    continue

                # Try to extract image bytes
                img_bytes = _extract_pdf_image_at_bbox(
                    doc, page, page_images, bbox, xref_to_bytes
                )

                if img_bytes:
                    blocks.append(ContentBlock(
                        type="image",
                        image_bytes=img_bytes,
                        image_mime="image/png",
                        width_pt=width_pt,
                        height_pt=height_pt,
                        image_index=image_counter,
                    ))
                    image_counter += 1

    doc.close()
    return _consolidate_text_blocks(blocks)


def _extract_pdf_image_at_bbox(
    doc, page, page_images, bbox, xref_cache: dict[int, bytes]
) -> bytes | None:
    """Extract image bytes from a PDF page at a given bounding box."""
    # Try to match with get_images() results by position
    for img_info in page_images:
        xref = img_info[0]

        if xref in xref_cache:
            return xref_cache[xref]

        try:
            pix = fitz.Pixmap(doc, xref)
            if pix.n > 4:  # CMYK -> RGB
                pix = fitz.Pixmap(fitz.csRGB, pix)
            png_bytes = pix.tobytes("png")
            xref_cache[xref] = png_bytes
            return png_bytes
        except Exception:
            continue

    # Fallback: clip the page region as an image
    try:
        clip = fitz.Rect(bbox)
        mat = fitz.Matrix(2, 2)  # 2x zoom for quality
        pix = page.get_pixmap(matrix=mat, clip=clip)
        return pix.tobytes("png")
    except Exception:
        return None


# ---------------------------------------------------------------------------
# Marked text building
# ---------------------------------------------------------------------------

def build_marked_text(blocks: list[ContentBlock]) -> str:
    """Build text with {{FIGURE_ORIGINALE_N}} markers at image positions."""
    parts = []
    for block in blocks:
        if block.type == "text":
            parts.append(block.text)
        elif block.type == "image":
            parts.append(f"{{{{FIGURE_ORIGINALE_{block.image_index + 1}}}}}")
    return "\n\n".join(parts)


# ---------------------------------------------------------------------------
# Adapted text splitting
# ---------------------------------------------------------------------------

def split_adapted_text(adapted_text: str, image_count: int) -> list[str]:
    """Split adapted text at {{FIGURE_ORIGINALE_N}} markers.

    Returns N+1 text segments (text before first image, between images, after last).
    """
    pattern = r"\{\{FIGURE_ORIGINALE_\d+\}\}"
    segments = re.split(pattern, adapted_text)
    return segments


# ---------------------------------------------------------------------------
# DOCX reconstruction
# ---------------------------------------------------------------------------

def rebuild_docx(blocks: list[ContentBlock], adapted_markdown: str) -> bytes:
    """Reconstruct a .docx with adapted text + original images."""
    doc = DocxDocument()

    image_blocks = [b for b in blocks if b.type == "image"]
    segments = split_adapted_text(adapted_markdown, len(image_blocks))

    image_idx = 0
    for seg_idx, segment in enumerate(segments):
        # Add text paragraphs
        if segment.strip():
            _add_markdown_to_docx(doc, segment.strip())

        # Add the corresponding image
        if image_idx < len(image_blocks):
            img_block = image_blocks[image_idx]
            img_stream = io.BytesIO(img_block.image_bytes)
            try:
                if img_block.width_pt and img_block.width_pt > 0:
                    # Cap at 6 inches (standard page width with margins)
                    width = min(Pt(img_block.width_pt), Inches(6))
                    doc.add_picture(img_stream, width=width)
                else:
                    doc.add_picture(img_stream, width=Inches(5))
            except Exception:
                doc.add_paragraph("[Image non insérée — format non supporté]")
            image_idx += 1

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_markdown_to_docx(doc: DocxDocument, text: str):
    """Add markdown text to a docx document with basic formatting."""
    lines = text.split("\n")

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # Headings
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("---"):
            # Horizontal rule -> empty paragraph
            doc.add_paragraph("—" * 40)
        elif stripped.startswith("> "):
            # Blockquote
            p = doc.add_paragraph(stripped[2:])
            p.style = "Quote" if "Quote" in [s.name for s in doc.styles] else None
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            # Numbered list
            text_after = re.sub(r"^\d+\.\s", "", stripped)
            doc.add_paragraph(text_after, style="List Number")
        else:
            # Regular paragraph with bold handling
            p = doc.add_paragraph()
            _add_formatted_runs(p, stripped)


def _add_formatted_runs(paragraph, text: str):
    """Add text with **bold** formatting to a paragraph."""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# HTML image reinsertion
# ---------------------------------------------------------------------------

def rebuild_html_with_images(html: str, image_blocks: list[ContentBlock]) -> str:
    """Replace {{FIGURE_ORIGINALE_N}} markers in HTML with embedded base64 images."""
    result = html
    for img in image_blocks:
        marker = f"{{{{FIGURE_ORIGINALE_{img.image_index + 1}}}}}"
        b64 = base64.b64encode(img.image_bytes).decode("ascii")
        img_tag = (
            f'<div class="adapted-figure original-figure">'
            f'<img src="data:{img.image_mime};base64,{b64}" '
            f'style="max-width:100%;height:auto;" '
            f'alt="Figure originale {img.image_index + 1}" />'
            f"</div>"
        )
        result = result.replace(marker, img_tag)
        # Also handle HTML-escaped braces (markdown may escape them)
        result = result.replace(
            f"%7B%7BFIGURE_ORIGINALE_{img.image_index + 1}%7D%7D", img_tag
        )
    return result


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _consolidate_text_blocks(blocks: list[ContentBlock]) -> list[ContentBlock]:
    """Merge consecutive text blocks into single blocks."""
    if not blocks:
        return blocks

    consolidated: list[ContentBlock] = []
    for block in blocks:
        if (
            block.type == "text"
            and consolidated
            and consolidated[-1].type == "text"
        ):
            consolidated[-1].text += "\n\n" + block.text
        else:
            consolidated.append(block)

    return consolidated
